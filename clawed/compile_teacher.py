"""Teacher-view DOCX compiler for MasterContent.

Compiles a MasterContent object into a teacher-facing Word document with
full answer keys, teacher scripts, and all instructional notes.
No LLM calls — pure mechanical compilation.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from clawed.master_content import MasterContent

logger = logging.getLogger(__name__)


async def compile_teacher_view(
    master: "MasterContent",
    images: dict[str, Path],
    output_dir: Path,
) -> Path:
    """Compile a teacher-facing DOCX from a MasterContent object.

    Includes full answer keys, teacher scripts (italicised), guided notes
    with answers filled in, station answer keys, and differentiation notes.

    Args:
        master: The MasterContent source-of-truth object.
        images: Mapping of image_spec strings to local file Paths.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    from clawed.io import safe_filename

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── helpers ──────────────────────────────────────────────────────

    def _heading(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def _para(text: str, bold: bool = False, italic: bool = False,
               size_pt: int = 11, color: tuple[int, int, int] | None = None) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _embed_image(image_spec: str, width_inches: float = 4.5) -> None:
        """Embed a pre-fetched image if its spec is in the images dict."""
        if not image_spec:
            return
        path = images.get(image_spec)
        if path and Path(path).exists():
            try:
                doc.add_picture(str(path), width=Inches(width_inches))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                logger.debug("Could not embed image %r: %s", image_spec, exc)

    def _shaded_cell(cell, fill_hex: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill_hex,
        })
        tcPr.append(shd)

    # ── Title / metadata header ───────────────────────────────────────

    title_para = doc.add_heading(master.title, level=0)

    meta_lines = [
        f"Subject: {master.subject}  |  Grade: {master.grade_level}  |  "
        f"Duration: {master.duration_minutes} min",
        f"Topic: {master.topic}",
        f"Objective: {master.objective}",
    ]
    for line in meta_lines:
        _para(line)

    if master.standards:
        _para("Standards: " + ", ".join(master.standards), bold=True)

    if master.materials_needed:
        _para("Materials: " + ", ".join(master.materials_needed))

    doc.add_paragraph("")

    # ── Vocabulary ────────────────────────────────────────────────────

    if master.vocabulary:
        _heading("Vocabulary")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for cell, label in zip(hdr_cells, ["Term", "Definition", "Context Sentence"]):
            _shaded_cell(cell, "BDD7EE")
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True

        for entry in master.vocabulary:
            row = table.add_row().cells
            row[0].text = entry.term
            row[1].text = entry.definition
            row[2].text = entry.context_sentence
            if entry.image_spec:
                _embed_image(entry.image_spec, width_inches=1.5)

        doc.add_paragraph("")

    # ── Do Now ────────────────────────────────────────────────────────

    _heading("Do Now")
    _para(master.do_now.stimulus)
    if master.do_now.questions:
        for i, q in enumerate(master.do_now.questions, 1):
            _para(f"{i}. {q}")
    # Teacher answer key
    if master.do_now.answers:
        _para("ANSWERS:", bold=True, color=(0x00, 0x70, 0xC0))
        for i, ans in enumerate(master.do_now.answers, 1):
            _para(f"{i}. {ans}", italic=True)
    doc.add_paragraph("")

    # ── Direct Instruction ────────────────────────────────────────────

    _heading("Direct Instruction")
    for section in master.direct_instruction:
        _heading(section.heading, level=2)
        _para(section.content)
        if section.key_points:
            _para("Key Points:", bold=True)
            for kp in section.key_points:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(kp)
        # Teacher script in italics
        if section.teacher_script:
            _para("Teacher Script:", bold=True, color=(0x70, 0x30, 0xA0))
            _para(section.teacher_script, italic=True, color=(0x70, 0x30, 0xA0))
        _embed_image(section.image_spec)
        doc.add_paragraph("")

    # ── Guided Notes (answers filled in) ─────────────────────────────

    if master.guided_notes:
        _heading("Guided Notes (Answer Key)")
        for note in master.guided_notes:
            _para(f"Prompt: {note.prompt}", bold=True)
            _para(f"Answer: {note.answer}", italic=True, color=(0x00, 0x70, 0xC0))
            if note.section_ref:
                _para(f"(Ref: {note.section_ref})", size_pt=9, color=(0x66, 0x66, 0x66))
            doc.add_paragraph("")

    # ── Primary Sources ────────────────────────────────────────────────

    if master.primary_sources:
        _heading("Primary Sources")
        for ps in master.primary_sources:
            _heading(ps.title, level=2)
            _para(f"Type: {ps.source_type}  |  Attribution: {ps.attribution}")
            _para(ps.content_text)
            if ps.scaffolding_questions:
                _para("Scaffolding Questions:", bold=True)
                for sq in ps.scaffolding_questions:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(sq)
            _embed_image(ps.image_spec)
            doc.add_paragraph("")

    # ── Stations (with answer keys) ───────────────────────────────────

    if master.stations:
        _heading("Learning Stations")
        for station in master.stations:
            _heading(station.title, level=2)
            _para(f"Source: {station.source_ref}")
            _para(f"Task: {station.task}")
            _para("Student Directions:", bold=True)
            _para(station.student_directions)
            _para("Answer Key:", bold=True, color=(0xC0, 0x00, 0x00))
            _para(station.teacher_answer_key, italic=True, color=(0xC0, 0x00, 0x00))
            doc.add_paragraph("")

    # ── Exit Ticket (with answers) ────────────────────────────────────

    if master.exit_ticket:
        _heading("Exit Ticket")
        for i, sq in enumerate(master.exit_ticket, 1):
            _para(f"Q{i} Stimulus ({sq.stimulus_type}): {sq.stimulus}", bold=True)
            if sq.stimulus_image_spec:
                _embed_image(sq.stimulus_image_spec)
            _para(f"Question: {sq.question}")
            _para(f"Expected Answer: {sq.answer}", italic=True, color=(0xC0, 0x00, 0x00))
            if sq.cognitive_level:
                _para(f"Cognitive Level: {sq.cognitive_level}", size_pt=9, color=(0x66, 0x66, 0x66))
            doc.add_paragraph("")

    # ── Differentiation ───────────────────────────────────────────────

    diff = master.differentiation
    _heading("Differentiation")
    if diff.struggling:
        _para("Struggling Learners:", bold=True)
        for item in diff.struggling:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    if diff.advanced:
        _para("Advanced Learners:", bold=True)
        for item in diff.advanced:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    if diff.ell:
        _para("ELL Support:", bold=True)
        for item in diff.ell:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    doc.add_paragraph("")

    # ── Independent Work ──────────────────────────────────────────────

    if master.independent_work:
        _heading("Independent Work")
        _para(master.independent_work.task)
        if master.independent_work.rubric_snippet:
            _para("Rubric:", bold=True)
            _para(master.independent_work.rubric_snippet)
        if master.independent_work.exemplar:
            _para("Exemplar:", bold=True)
            _para(master.independent_work.exemplar)
        doc.add_paragraph("")

    # ── Homework ──────────────────────────────────────────────────────

    if master.homework:
        _heading("Homework")
        _para(master.homework)

    # ── Save ──────────────────────────────────────────────────────────

    safe = safe_filename(master.title)
    out_path = output_dir / f"{safe}_teacher.docx"
    doc.save(str(out_path))
    logger.info("Teacher view saved to %s", out_path)
    return out_path
