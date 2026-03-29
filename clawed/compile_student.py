"""Student-view DOCX compiler for MasterContent.

Compiles a MasterContent object into a student-facing handout Word document.
Answer keys, teacher scripts, and station answer keys are omitted; guided
notes show prompts with blank lines instead of answers.
No LLM calls — pure mechanical compilation.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from clawed.master_content import MasterContent

logger = logging.getLogger(__name__)

_BLANK = "_____________"


async def compile_student_view(
    master: "MasterContent",
    images: dict[str, Path],
    output_dir: Path,
) -> Path:
    """Compile a student-facing DOCX from a MasterContent object.

    Guided notes show prompts with blank lines (not answers).  Teacher scripts,
    station answer keys, exit-ticket answers, and differentiation notes are all
    omitted — this is the print-ready student handout.

    Args:
        master: The MasterContent source-of-truth object.
        images: Mapping of image_spec strings to local file Paths.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    from clawed.io import safe_filename

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── helpers ──────────────────────────────────────────────────────

    def _heading(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def _para(text: str, bold: bool = False, italic: bool = False,
               size_pt: int = 11, color: tuple[int, int, int] | None = None) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _embed_image(image_spec: str, width_inches: float = 4.5) -> None:
        if not image_spec:
            return
        path = images.get(image_spec)
        if path and Path(path).exists():
            try:
                doc.add_picture(str(path), width=Inches(width_inches))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                logger.debug("Could not embed image %r: %s", image_spec, exc)

    def _shaded_cell(cell, fill_hex: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()  # noqa: N806
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill_hex,
        })
        tcPr.append(shd)

    # ── Title / metadata header ───────────────────────────────────────

    doc.add_heading(master.title, level=0)

    meta_lines = [
        f"Subject: {master.subject}  |  Grade: {master.grade_level}",
        f"Topic: {master.topic}",
        f"Objective: {master.objective}",
    ]
    for line in meta_lines:
        _para(line)

    doc.add_paragraph("Name: ___________________________  Date: ___________  Period: _____")
    doc.add_paragraph("")

    # ── Vocabulary ────────────────────────────────────────────────────

    if master.vocabulary:
        _heading("Vocabulary")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for cell, label in zip(hdr_cells, ["Term", "Definition", "Context Sentence"]):
            _shaded_cell(cell, "BDD7EE")
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True

        for entry in master.vocabulary:
            row = table.add_row().cells
            row[0].text = entry.term
            row[1].text = entry.definition
            row[2].text = entry.context_sentence
            if entry.image_spec:
                _embed_image(entry.image_spec, width_inches=1.5)

        doc.add_paragraph("")

    # ── Do Now ────────────────────────────────────────────────────────

    _heading("Do Now")
    _para(master.do_now.stimulus)
    if master.do_now.questions:
        for i, q in enumerate(master.do_now.questions, 1):
            _para(f"{i}. {q}")
            _para(_BLANK)
            _para("")
    doc.add_paragraph("")

    # ── Direct Instruction ────────────────────────────────────────────
    # No teacher script in student view

    _heading("Direct Instruction")
    for section in master.direct_instruction:
        _heading(section.heading, level=2)
        _para(section.content)
        if section.key_points:
            _para("Key Points:", bold=True)
            for kp in section.key_points:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(kp)
        _embed_image(section.image_spec)
        doc.add_paragraph("")

    # ── Guided Notes (blanks — no answers) ───────────────────────────

    if master.guided_notes:
        _heading("Guided Notes")
        for note in master.guided_notes:
            _para(note.prompt, bold=True)
            _para(_BLANK)
            doc.add_paragraph("")

    # ── Primary Sources ────────────────────────────────────────────────

    if master.primary_sources:
        _heading("Primary Sources")
        for ps in master.primary_sources:
            _heading(ps.title, level=2)
            _para(f"Type: {ps.source_type}  |  Attribution: {ps.attribution}")
            _para(ps.content_text)
            if ps.scaffolding_questions:
                _para("Questions:", bold=True)
                for sq in ps.scaffolding_questions:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(sq)
                    doc.add_paragraph(_BLANK)
            _embed_image(ps.image_spec)
            doc.add_paragraph("")

    # ── Stations (student directions only — no answer key) ────────────

    if master.stations:
        _heading("Learning Stations")
        for station in master.stations:
            _heading(station.title, level=2)
            _para(f"Task: {station.task}")
            _para("Directions:", bold=True)
            _para(station.student_directions)
            _para(_BLANK)
            doc.add_paragraph("")

    # ── Exit Ticket (stimulus + question, no answer) ──────────────────

    if master.exit_ticket:
        _heading("Exit Ticket")
        for i, sq in enumerate(master.exit_ticket, 1):
            _para(f"Q{i}: {sq.stimulus}", bold=True)
            if sq.stimulus_image_spec:
                _embed_image(sq.stimulus_image_spec)
            _para(sq.question)
            _para(_BLANK)
            doc.add_paragraph("")

    # ── Differentiation (visible to students for self-awareness) ──────
    # Per spec: differentiation section is included for student view too

    diff = master.differentiation
    _heading("Support Strategies")
    if diff.struggling:
        _para("If you need extra support:", bold=True)
        for item in diff.struggling:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    if diff.advanced:
        _para("Challenge extension:", bold=True)
        for item in diff.advanced:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    if diff.ell:
        _para("Language support:", bold=True)
        for item in diff.ell:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    doc.add_paragraph("")

    # ── Homework ──────────────────────────────────────────────────────

    if master.homework:
        _heading("Homework")
        _para(master.homework)

    # ── Save ──────────────────────────────────────────────────────────

    safe = safe_filename(master.title)
    out_path = output_dir / f"{safe}_student.docx"
    doc.save(str(out_path))
    logger.info("Student view saved to %s", out_path)
    return out_path
