"""Student packet export — the document students actually work through in class.

Produces a 4-6 page DOCX workbook with:
- Do Now prompt with response lines
- Key vocabulary with definitions
- Fill-in-the-blank guided notes
- Station sections with full primary source text, context, and analysis questions
- Graphic organizer table
- Exit ticket with sentence starters

Consumes a StudentPacket model (v2.3+) or falls back to legacy dict format.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def export_student_packet_docx(
    packet: Any,
    subject: str = "",
    output_dir: Path | None = None,
    agent_name: str = "Claw-ED",
) -> Path:
    """Export a StudentPacket to a professionally formatted DOCX workbook."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    from clawed.export_theme import get_color_theme
    from clawed.sanitize import sanitize_text

    theme = get_color_theme(subject)
    primary_hex = theme["primary"]
    primary_rgb = RGBColor(
        int(primary_hex[:2], 16), int(primary_hex[2:4], 16), int(primary_hex[4:6], 16),
    )
    accent_hex = theme.get("accent", primary_hex)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Helper functions ──────────────────────────────────────────────

    def _get(field, default=""):
        if hasattr(packet, field):
            return getattr(packet, field, default)
        if isinstance(packet, dict):
            return packet.get(field, default)
        return default

    def _section_heading(text: str) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(sanitize_text(text).upper())
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = primary_rgb
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {qn("w:val"): "single", qn("w:sz"): "8", qn("w:space"): "1", qn("w:color"): primary_hex},
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_lines(count: int = 3) -> None:
        for _ in range(count):
            p = doc.add_paragraph("")
            p.paragraph_format.space_after = Pt(12)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "BBBBBB"},
            )
            pBdr.append(bottom)
            pPr.append(pBdr)

    def _add_table_borders(table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
        borders = tblPr.makeelement(qn("w:tblBorders"), {})
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = borders.makeelement(
                qn(f"w:{edge}"),
                {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "0", qn("w:color"): "444444"},
            )
            borders.append(el)
        tblPr.append(borders)

    def _shade_cell(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(
            qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color},
        )
        tcPr.append(shading)

    # ── Title + Header ────────────────────────────────────────────────
    title = sanitize_text(_get("title", "Student Packet"))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = primary_rgb

    # Aim
    aim = sanitize_text(_get("aim", ""))
    if aim:
        aim_para = doc.add_paragraph()
        aim_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        aim_para.paragraph_format.space_after = Pt(4)
        r = aim_para.add_run(f"Aim: {aim}")
        r.bold = True
        r.font.size = Pt(12)

    # Name/Date/Period
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run("Name: _________________________   Date: ______________   Period: ________")
    meta_run.font.size = Pt(10)

    # ── Do Now ────────────────────────────────────────────────────────
    do_now = sanitize_text(_get("do_now_prompt") or _get("do_now", ""))
    if do_now:
        _section_heading("Do Now")
        p = doc.add_paragraph(do_now)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)
        lines = _get("do_now_response_lines", 4)
        if isinstance(lines, int):
            _add_lines(lines)
        else:
            _add_lines(4)

    # ── Key Vocabulary ────────────────────────────────────────────────
    vocab = _get("vocabulary", [])
    if vocab:
        _section_heading("Key Vocabulary")
        for item in vocab:
            if isinstance(item, dict):
                term = sanitize_text(item.get("term", ""))
                defn = sanitize_text(item.get("definition", ""))
            else:
                term = sanitize_text(getattr(item, "term", ""))
                defn = sanitize_text(getattr(item, "definition", ""))
            if term:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"{term}: ")
                r.bold = True
                r.font.size = Pt(11)
                p.add_run(defn)

    # ── Guided Notes (fill-in-the-blank) ──────────────────────────────
    guided = _get("guided_notes", [])
    if guided:
        _section_heading("Guided Notes")
        p_intro = doc.add_paragraph("Directions: Fill in the blanks as we go through the lesson.")
        p_intro.italic = True
        p_intro.paragraph_format.space_after = Pt(6)
        for i, item in enumerate(guided, 1):
            if isinstance(item, dict):
                sentence = sanitize_text(item.get("sentence_with_blank", ""))
            else:
                sentence = sanitize_text(getattr(item, "sentence_with_blank", ""))
            if sentence:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(f"{i}. {sentence}")
                r.font.size = Pt(11)

    # ── Stations / Primary Source Documents ────────────────────────────
    stations = _get("stations", []) or _get("source_excerpts", [])
    if stations:
        _section_heading("Document Analysis")
        for station in stations:
            if isinstance(station, dict):
                label = sanitize_text(station.get("document_label", station.get("title", "")))
                ctx = sanitize_text(station.get("context", ""))
                full_text = sanitize_text(station.get("full_text", station.get("text", "")))
                author = sanitize_text(station.get("author", station.get("attribution", "")))
                date = sanitize_text(station.get("date", ""))
                questions = station.get("analysis_questions", [])
            else:
                label = sanitize_text(getattr(station, "document_label", "") or getattr(station, "title", ""))
                ctx = sanitize_text(getattr(station, "context", ""))
                full_text = sanitize_text(getattr(station, "full_text", ""))
                author = sanitize_text(getattr(station, "author", ""))
                date = sanitize_text(getattr(station, "date", ""))
                questions = getattr(station, "analysis_questions", [])

            # Station header
            if label:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(label)
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = primary_rgb

            # Context paragraph
            if ctx:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run("Context: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(ctx)
                r2.font.size = Pt(10)
                r2.italic = True

            # Full source text — indented, quoted
            if full_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(f'\u201c{full_text}\u201d')
                r.italic = True
                r.font.size = Pt(11)
                # Attribution
                if author or date:
                    attrib_parts = [x for x in [author, date] if x]
                    a = doc.add_paragraph()
                    a.paragraph_format.left_indent = Inches(0.4)
                    a.paragraph_format.space_after = Pt(6)
                    ar = a.add_run(f"\u2014 {', '.join(attrib_parts)}")
                    ar.font.size = Pt(9)
                    ar.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Analysis questions with response lines
            if questions:
                for qi, q in enumerate(questions, 1):
                    q_text = sanitize_text(q) if isinstance(q, str) else sanitize_text(str(q))
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    r = p.add_run(f"{qi}. {q_text}")
                    r.font.size = Pt(11)
                    _add_lines(5)

    # ── Graphic Organizer ─────────────────────────────────────────────
    organizer = _get("graphic_organizer")
    if organizer:
        if isinstance(organizer, dict):
            org_title = sanitize_text(organizer.get("title", "Graphic Organizer"))
            org_instructions = sanitize_text(organizer.get("instructions", ""))
            columns = organizer.get("columns", [])
            num_rows = organizer.get("num_rows", 4)
        else:
            org_title = sanitize_text(getattr(organizer, "title", "Graphic Organizer"))
            org_instructions = sanitize_text(getattr(organizer, "instructions", ""))
            columns = getattr(organizer, "columns", [])
            num_rows = getattr(organizer, "num_rows", 4)

        if columns:
            _section_heading(org_title)
            if org_instructions:
                p = doc.add_paragraph(org_instructions)
                p.italic = True
                p.paragraph_format.space_after = Pt(6)

            table = doc.add_table(rows=num_rows + 1, cols=len(columns))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row with shading
            for i, col in enumerate(columns):
                cell = table.rows[0].cells[i]
                cell.text = sanitize_text(col) if isinstance(col, str) else str(col)
                _shade_cell(cell, accent_hex)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(10)
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Empty data rows with minimum height for writing
            for row_idx in range(1, num_rows + 1):
                for cell in table.rows[row_idx].cells:
                    cell.text = ""
                    # Add space for writing
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(30)

            _add_table_borders(table)

    # ── Exit Ticket ───────────────────────────────────────────────────
    questions = _get("exit_ticket_questions", [])
    if questions:
        _section_heading("Exit Ticket")
        for i, q in enumerate(questions, 1):
            q_text = sanitize_text(q) if isinstance(q, str) else sanitize_text(str(q))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run(f"{i}. {q_text}")
            r.bold = True
            r.font.size = Pt(11)
            _add_lines(5)

    # Sentence starters
    starters = _get("sentence_starters", [])
    if starters:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run("Sentence Starters to Help You:")
        r.bold = True
        r.italic = True
        r.font.size = Pt(10)
        for starter in starters:
            s_text = sanitize_text(starter) if isinstance(starter, str) else str(starter)
            bp = doc.add_paragraph(style="List Bullet")
            br = bp.add_run(s_text)
            br.italic = True
            br.font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"Created with {agent_name}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──────────────────────────────────────────────────────────
    if output_dir is None:
        output_dir = Path("clawed_output").resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_title = "".join(c for c in title[:40] if c.isalnum() or c in " _-").strip() or "packet"
    out_path = output_dir / f"{safe_title}_packet.docx"
    doc.save(str(out_path))
    return out_path


# ── Backward-compatible alias ─────────────────────────────────────────


def export_handout_docx(
    handout_data: dict[str, Any],
    subject: str = "",
    output_dir: Path | None = None,
) -> Path:
    """Legacy alias — accepts old-format handout dict, routes to new exporter."""
    return export_student_packet_docx(handout_data, subject=subject, output_dir=output_dir)
