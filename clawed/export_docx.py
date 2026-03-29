"""Word document (DOCX) export for lesson plans and student handouts.

Generates full lesson plan documents and print-ready student worksheets.
"""

from __future__ import annotations

import asyncio
import logging
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING, Any

from clawed.export_theme import _resolve_output, get_color_theme

if TYPE_CHECKING:
    from clawed.models import DailyLesson, TeacherPersona

logger = logging.getLogger(__name__)


def _run_async_safe(coro):
    """Run an async coroutine, handling both sync and async calling contexts.

    When called from inside a running event loop (e.g., agent_core tools),
    uses a thread to avoid nested asyncio.run() errors.  When called from
    plain sync code (CLI), uses asyncio.run() directly.
    """
    try:
        asyncio.get_running_loop()
        # We're inside an event loop — run in a worker thread
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            return pool.submit(asyncio.run, coro).result(timeout=30)
    except RuntimeError:
        # No running loop — safe to use asyncio.run()
        return asyncio.run(coro)


# ── Image helpers (DOCX-specific) ─────────────────────────────────────


def _docx_add_image(
    doc: Any,  # docx.Document
    topic: str,
    subject: str,
    width_inches: float = 5.0,
    caption: str = "",
) -> bool:
    """Try to fetch and embed an academic image into the DOCX.

    Fails silently -- handout works fine without images.
    Returns True if an image was successfully added.
    """
    try:
        from docx.shared import Inches, Pt

        from clawed.slide_images import fetch_slide_image

        img_path = _run_async_safe(fetch_slide_image(topic, subject=subject))
        if img_path and img_path.exists():
            doc.add_picture(str(img_path), width=Inches(width_inches))
            # Center the image
            last_para = doc.paragraphs[-1]
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add caption if provided
            if caption:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.font.size = Pt(9)
                cap_run.font.italic = True
                cap_run.font.name = "Calibri"
                from docx.shared import RGBColor
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            return True
    except Exception:
        pass  # Images are a bonus, not a requirement
    return False


def _docx_add_content_image(
    doc: Any,  # docx.Document
    content_text: str,
    fallback_topic: str,
    subject: str,
    width_inches: float = 3.0,
) -> bool:
    """Fetch and embed an image based on content text, with auto-caption.

    Uses ``_extract_key_concepts`` to find the best search query, then
    adds a captioned image.  Returns True if an image was added.
    """
    try:
        from docx.shared import Inches, Pt

        from clawed.slide_images import _extract_key_concepts, fetch_content_image

        img_path = _run_async_safe(
            fetch_content_image(
                content_text,
                subject=subject,
                fallback_topic=fallback_topic,
            )
        )
        if img_path and img_path.exists():
            doc.add_picture(str(img_path), width=Inches(width_inches))
            last_para = doc.paragraphs[-1]
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Build caption from key concepts
            concepts = _extract_key_concepts(content_text)
            caption = ", ".join(concepts[:2]) if concepts else ""
            if caption:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.font.size = Pt(9)
                cap_run.font.italic = True
                cap_run.font.name = "Calibri"
                from docx.shared import RGBColor
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            return True
    except Exception:
        pass
    return False


# ── Callout box helper ────────────────────────────────────────────────


def _add_callout_box(doc, label, text, bg_hex, border_hex):
    """Add a colored callout box for differentiation notes."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell: narrow colored label
    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(1.2)
    left_cell.text = ""
    lp = left_cell.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.name = "Calibri"
    # Set cell background
    tc = left_cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): border_hex,
    })
    tcPr.append(shading)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Right cell: content with light background
    right_cell = table.rows[0].cells[1]
    right_cell.text = ""
    rp = right_cell.paragraphs[0]
    rr = rp.add_run(text)
    rr.font.size = Pt(10)
    rr.font.name = "Calibri"
    # Light background
    tc2 = right_cell._tc
    tcPr2 = tc2.get_or_add_tcPr()
    shading2 = tcPr2.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): bg_hex,
    })
    tcPr2.append(shading2)

    # Add spacing after table
    doc.add_paragraph("")


# ── Main lesson DOCX export ──────────────────────────────────────────


def export_lesson_docx(
    lesson: "DailyLesson",
    persona: "TeacherPersona",
    output_dir: Path | None = None,
    agent_name: str = "Claw-ED",
    admin_plan: Any = None,
) -> Path:
    """Generate a Word document from a lesson plan with embedded academic images.

    If an AdminLessonPlan is provided, generates an observation-ready document
    with per-section teacher/student actions, observer look-fors, anticipated
    responses, and teacher content knowledge.

    Returns the path to the saved .docx file.
    """
    # If we have an admin plan, use the enriched export
    if admin_plan is not None:
        return _export_admin_lesson_docx(lesson, persona, admin_plan, output_dir, agent_name)

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    from clawed.sanitize import sanitize_text

    doc = Document()

    # ── Subject color theme ──────────────────────────────────────────
    theme = get_color_theme(persona.subject_area or "")

    def _theme_rgb(key: str) -> "RGBColor":
        h = theme[key]
        return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    # ── Resolve teacher display name ─────────────────────────────────
    teacher_display_name = ""
    if persona and persona.name and persona.name != "My Teaching Persona":
        teacher_display_name = persona.name
    else:
        try:
            from clawed.models import AppConfig as _AppConfig
            _cfg = _AppConfig.load()
            if _cfg.teacher_profile and _cfg.teacher_profile.name:
                teacher_display_name = _cfg.teacher_profile.name
        except Exception:
            pass
    if not teacher_display_name:
        teacher_display_name = "Teacher"

    # ── Professional header ──────────────────────────────────────────
    header_section = doc.sections[0]
    header = header_section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(
        f"{teacher_display_name}  |  "
        f"{persona.subject_area or 'Education'}  |  "
        f"{date.today().strftime('%B %d, %Y')}"
    )
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Apply accent color as header bar background
    try:
        from docx.oxml.ns import qn as _qn
        _shd = header_para._p.get_or_add_pPr().makeelement(
            _qn("w:shd"),
            {
                _qn("w:val"): "clear",
                _qn("w:color"): "auto",
                _qn("w:fill"): theme["accent"],
            },
        )
        header_para._p.get_or_add_pPr().append(_shd)
    except Exception:
        pass

    # ── Professional footer ──────────────────────────────────────────
    footer = header_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"{agent_name}  |  Page")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Sanitize all text content before writing to the document
    _s_title = sanitize_text(lesson.title)
    _s_objective = sanitize_text(lesson.objective)
    _s_do_now = sanitize_text(lesson.do_now) if lesson.do_now else ""
    _s_direct = sanitize_text(lesson.direct_instruction) if lesson.direct_instruction else ""
    _s_guided = sanitize_text(lesson.guided_practice) if lesson.guided_practice else ""
    _s_independent = sanitize_text(lesson.independent_work) if lesson.independent_work else ""
    _s_homework = sanitize_text(lesson.homework) if lesson.homework else ""

    # Resolve subject for image searches
    subject = persona.subject_area or ""

    # Title
    doc.add_heading(_s_title, level=0)
    doc.add_paragraph(
        f"Teacher: {teacher_display_name}  |  "
        f"Lesson {lesson.lesson_number}  |  "
        f"{date.today().strftime('%B %d, %Y')}"
    )

    # Try to add a header image relevant to the lesson content
    _docx_add_content_image(
        doc,
        content_text=_s_title + " " + _s_objective,
        fallback_topic=_s_title,
        subject=subject,
        width_inches=5.5,
    )

    # Standards table — fallback to generic alignment if empty
    if not lesson.standards:
        try:
            from clawed.models import AppConfig as _AppConfig
            cfg = _AppConfig.load()
            if cfg.teacher_profile.state:
                lesson.standards = [
                    f"Aligned to {cfg.teacher_profile.state} "
                    f"{persona.subject_area or 'education'} standards"
                ]
        except Exception:
            pass

    if lesson.standards:
        doc.add_heading("Standards Addressed", level=2)
        table = doc.add_table(rows=1, cols=1)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Standard"
        for std in lesson.standards:
            row = table.add_row().cells
            row[0].text = std

    # Objective
    doc.add_heading("Objective (SWBAT)", level=2)
    doc.add_paragraph(_s_objective)

    # Materials
    if lesson.materials_needed:
        doc.add_heading("Materials Needed", level=2)
        for m in lesson.materials_needed:
            doc.add_paragraph(m, style="List Bullet")

    # Lesson-at-a-Glance timing table
    doc.add_heading("Lesson at a Glance", level=2)
    timing_table = doc.add_table(rows=1, cols=2)
    timing_table.style = "Light Grid Accent 1"
    hdr = timing_table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Time"

    timing_sections = [
        ("Do Now / Warm-Up", lesson.time_estimates.get("do_now", 5)),
        ("Direct Instruction", lesson.time_estimates.get("direct_instruction", 20)),
        ("Guided Practice", lesson.time_estimates.get("guided_practice", 15)),
        ("Independent Work", lesson.time_estimates.get("independent_work", 10)),
        ("Exit Ticket", 5),
    ]
    total = 0
    for section_name, minutes in timing_sections:
        if minutes:
            row = timing_table.add_row().cells
            row[0].text = section_name
            row[1].text = f"{minutes} min"
            total += minutes
    # Total row
    total_row = timing_table.add_row().cells
    total_row[0].text = "Total"
    total_row[1].text = f"{total} min"
    for cell in timing_table.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    # Lesson Sections — add relevant images to instruction sections
    sections = [
        ("Do Now / Warm-Up", _s_do_now, False),
        ("Direct Instruction", _s_direct, True),
        ("Guided Practice", _s_guided, False),
        ("Independent Work", _s_independent, False),
    ]
    for heading, content, add_img in sections:
        if content:
            time_key = heading.lower().replace(" / warm-up", "").replace(" ", "_")
            minutes = lesson.time_estimates.get(time_key, "")
            time_label = f" ({minutes} min)" if minutes else ""
            doc.add_heading(f"{heading}{time_label}", level=2)
            doc.add_paragraph(content)
            # Add a content-specific image to the direct instruction section
            if add_img:
                _docx_add_content_image(
                    doc,
                    content_text=content,
                    fallback_topic=_s_title,
                    subject=subject,
                    width_inches=4.0,
                )

    # Exit Ticket
    if lesson.exit_ticket:
        doc.add_heading("Exit Ticket", level=2)
        for i, q in enumerate(lesson.exit_ticket, 1):
            doc.add_paragraph(f"{i}. {sanitize_text(q.question)}")

    # Differentiation
    diff = lesson.differentiation
    if diff:
        doc.add_heading("Differentiation", level=2)
        if diff.struggling:
            text = ", ".join(diff.struggling) if isinstance(diff.struggling, list) else str(diff.struggling)
            _add_callout_box(doc, "Struggling\nLearners", sanitize_text(text), "FFF3CD", "D4A017")
        if diff.advanced:
            text = ", ".join(diff.advanced) if isinstance(diff.advanced, list) else str(diff.advanced)
            _add_callout_box(doc, "Advanced\nLearners", sanitize_text(text), "D1ECF1", "2B7A98")
        if diff.ell:
            text = ", ".join(diff.ell) if isinstance(diff.ell, list) else str(diff.ell)
            _add_callout_box(doc, "ELL\nSupport", sanitize_text(text), "D4EDDA", "2D8B4E")

    # Homework
    if _s_homework:
        doc.add_heading("Homework", level=2)
        doc.add_paragraph(_s_homework)

    # Footer
    doc.add_paragraph("")
    gen_footer = doc.add_paragraph(f"Generated by {agent_name}")
    gen_footer.runs[0].font.size = Pt(8)

    # ── Apply subject color theme to headings and shaded sections ─────
    try:
        from docx.oxml.ns import qn as _qn

        primary_rgb = _theme_rgb("primary")
        bg_light_hex = theme["bg_light"]
        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading 2"):
                # Apply primary theme color to Heading 2 runs
                for run in para.runs:
                    run.font.color.rgb = primary_rgb
            # Apply bg_light shading to vocabulary/standards table containers
            # (tables are handled separately below)

        # Apply bg_light to standards table cells (vocabulary boxes)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _cell_shd = cell._tc.get_or_add_tcPr().makeelement(
                        _qn("w:shd"),
                        {
                            _qn("w:val"): "clear",
                            _qn("w:color"): "auto",
                            _qn("w:fill"): bg_light_hex.lstrip("#"),
                        },
                    )
                    cell._tc.get_or_add_tcPr().append(_cell_shd)
    except Exception:
        pass  # Theming is cosmetic — don't break export

    out = _resolve_output(output_dir, lesson, ".docx")
    doc.save(str(out))
    return out


# ── Student handout export ────────────────────────────────────────────


def export_student_handout(
    lesson: "DailyLesson",
    persona: "TeacherPersona",
    output_dir: Path | None = None,
    agent_name: str = "Claw-ED",
) -> Path:
    """Generate a print-ready student handout (DOCX worksheet).

    Produces a clean, self-contained 1-2 page document suitable for
    printing on standard letter paper.  Includes:
    - Header with lesson title, teacher name, date
    - Bordered Do Now box with lined response area
    - Aim / Objective line
    - Core content section (direct instruction text)
    - Numbered activity questions with lined response areas
    - Exit ticket questions with response lines
    - Footer with Name / Date / Period blanks

    Returns the path to the saved .docx file.
    """
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    from clawed.sanitize import sanitize_text

    doc = Document()

    # Sanitize all lesson text fields for the handout
    lesson.title = sanitize_text(lesson.title)
    lesson.objective = sanitize_text(lesson.objective)
    lesson.do_now = sanitize_text(lesson.do_now) if lesson.do_now else ""
    lesson.direct_instruction = sanitize_text(lesson.direct_instruction) if lesson.direct_instruction else ""
    lesson.guided_practice = sanitize_text(lesson.guided_practice) if lesson.guided_practice else ""
    lesson.independent_work = sanitize_text(lesson.independent_work) if lesson.independent_work else ""
    for q in lesson.exit_ticket:
        q.question = sanitize_text(q.question)

    # ── Page setup ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Default font ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)

    # ── Header: Title, Teacher, Date, Period ──────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(lesson.title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Resolve teacher display name for handout
    _handout_teacher_name = ""
    if persona and persona.name and persona.name != "My Teaching Persona":
        _handout_teacher_name = persona.name
    else:
        try:
            from clawed.models import AppConfig as _AppConfig
            _cfg = _AppConfig.load()
            if _cfg.teacher_profile and _cfg.teacher_profile.name:
                _handout_teacher_name = _cfg.teacher_profile.name
        except Exception:
            pass
    if not _handout_teacher_name:
        _handout_teacher_name = "Teacher"

    meta_run = meta_para.add_run(
        f"{_handout_teacher_name}  |  {date.today().strftime('%B %d, %Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta_run.font.name = "Calibri"

    # ── Header image (below title, max 2 images total) ───────────────
    subject = persona.subject_area or ""
    _handout_image_count = 0
    if _handout_image_count < 2:
        if _docx_add_content_image(
            doc,
            content_text=lesson.title + " " + lesson.objective,
            fallback_topic=lesson.title,
            subject=subject,
            width_inches=3.0,
        ):
            _handout_image_count += 1

    # ── Do Now box ────────────────────────────────────────────────────
    if lesson.do_now:
        _handout_section_heading(doc, "Do Now")
        do_now_table = doc.add_table(rows=2, cols=1)
        do_now_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(do_now_table)

        # Prompt cell
        prompt_cell = do_now_table.rows[0].cells[0]
        prompt_cell.text = ""
        prompt_para = prompt_cell.paragraphs[0]
        prompt_run = prompt_para.add_run(lesson.do_now)
        prompt_run.font.size = Pt(11)
        prompt_run.font.name = "Calibri"

        # Response area with lines
        response_cell = do_now_table.rows[1].cells[0]
        response_cell.text = ""
        _add_lined_space(response_cell, line_count=4)

    # ── Aim / Objective ───────────────────────────────────────────────
    _handout_section_heading(doc, "Aim")
    aim_para = doc.add_paragraph(lesson.objective)
    aim_para.paragraph_format.space_after = Pt(4)

    # ── Core Content ──────────────────────────────────────────────────
    if lesson.direct_instruction:
        _handout_section_heading(doc, "Key Content")
        # Include a condensed version -- long DI sections get trimmed
        content_text = lesson.direct_instruction
        if len(content_text) > 1200:
            # Take first ~1200 chars at a sentence boundary
            cutoff = content_text[:1200].rfind(". ")
            if cutoff > 600:
                content_text = content_text[: cutoff + 1]
        content_para = doc.add_paragraph(content_text)
        content_para.paragraph_format.space_after = Pt(6)
        for run in content_para.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

        # Add content image next to direct instruction (max 2 total)
        if _handout_image_count < 2:
            if _docx_add_content_image(
                doc,
                content_text=lesson.direct_instruction,
                fallback_topic=lesson.title,
                subject=subject,
                width_inches=3.0,
            ):
                _handout_image_count += 1

    # ── Activity Section (Guided Practice) ────────────────────────────
    if lesson.guided_practice:
        _handout_section_heading(doc, "Activity")
        _add_numbered_content_with_lines(doc, lesson.guided_practice)

    # ── Independent Work ──────────────────────────────────────────────
    if lesson.independent_work:
        _handout_section_heading(doc, "Independent Practice")
        _add_numbered_content_with_lines(doc, lesson.independent_work)

    # ── Exit Ticket ───────────────────────────────────────────────────
    if lesson.exit_ticket:
        _handout_section_heading(doc, "Exit Ticket")
        exit_table = doc.add_table(rows=len(lesson.exit_ticket) * 2, cols=1)
        exit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(exit_table)

        for i, q in enumerate(lesson.exit_ticket):
            q_cell = exit_table.rows[i * 2].cells[0]
            q_cell.text = ""
            q_para = q_cell.paragraphs[0]
            q_run = q_para.add_run(f"{i + 1}. {q.question}")
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = "Calibri"

            ans_cell = exit_table.rows[i * 2 + 1].cells[0]
            ans_cell.text = ""
            _add_lined_space(ans_cell, line_count=3)

    # ── Footer: Name / Date / Period ──────────────────────────────────
    doc.add_paragraph("")  # spacer
    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(3.0)
    footer_table.columns[1].width = Inches(2.5)
    footer_table.columns[2].width = Inches(1.5)

    labels = ["Name: _______________", "Date: _______________", "Period: ______"]
    for idx, label in enumerate(labels):
        cell = footer_table.rows[0].cells[idx]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(label)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Remove borders from footer table
    _remove_table_borders(footer_table)

    # ── Watermark ─────────────────────────────────────────────────────
    wm_para = doc.add_paragraph()
    wm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wm_run = wm_para.add_run(f"Generated by {agent_name}")
    wm_run.font.size = Pt(8)
    wm_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──────────────────────────────────────────────────────────
    out = _resolve_output(output_dir, lesson, "_handout.docx")
    doc.save(str(out))
    return out


# ── DOCX formatting helpers ──────────────────────────────────────────


def _handout_section_heading(doc: Any, text: str) -> None:
    """Add a styled section heading for the student handout."""
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    # Add a bottom border via XML
    from docx.oxml.ns import qn

    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "1A365D",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_lined_space(cell: Any, line_count: int = 3) -> None:
    """Add blank lines with bottom borders to simulate writing lines."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for i in range(line_count):
        para = cell.add_paragraph("")
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
        # Add bottom border to simulate a writing line
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "CCCCCC",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)


def _add_numbered_content_with_lines(doc: Any, text: str) -> None:
    """Parse text into numbered items and add lined response areas.

    If the text contains numbered items (1. / 2. / etc.), each gets its own
    response area.  Otherwise the full text is shown with a single response area.
    """
    import re

    from docx.shared import Pt

    # Try to split on numbered items
    items = re.split(r"(?:^|\n)\s*(\d+)[.)]\s*", text)

    # items[0] is preamble, then alternating [number, content, number, content...]
    preamble = items[0].strip() if items else ""
    numbered: list[str] = []
    if len(items) > 2:
        for i in range(1, len(items) - 1, 2):
            numbered.append(items[i + 1].strip() if i + 1 < len(items) else "")

    if numbered:
        if preamble:
            p = doc.add_paragraph(preamble)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
        for idx, item in enumerate(numbered, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {item}")
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)
            # Add a small lined area after each item
            line_para = doc.add_paragraph("")
            line_para.paragraph_format.space_after = Pt(10)
            from docx.oxml.ns import qn

            pPr = line_para._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "1",
                    qn("w:color"): "CCCCCC",
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)
    else:
        # No numbered items — show the text with a response area
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        # Add 3 lined spaces
        for _ in range(3):
            line_para = doc.add_paragraph("")
            line_para.paragraph_format.space_after = Pt(10)
            from docx.oxml.ns import qn

            pPr = line_para._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "1",
                    qn("w:color"): "CCCCCC",
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)


def _set_table_borders(table: Any) -> None:
    """Set solid borders on all sides of a table."""
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "0",
                qn("w:color"): "444444",
            },
        )
        borders.append(element)
    tblPr.append(borders)


def _remove_table_borders(table: Any) -> None:
    """Remove all borders from a table."""
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"},
        )
        borders.append(element)
    tblPr.append(borders)


# ── Admin lesson plan export ──────────────────────────────────────────


def _export_admin_lesson_docx(
    lesson: "DailyLesson",
    persona: "TeacherPersona",
    admin_plan: Any,
    output_dir: Path | None = None,
    agent_name: str = "Claw-ED",
) -> Path:
    """Generate an observation-ready lesson plan DOCX with multi-column tables.

    This produces the format administrators expect: per-section teacher actions,
    student actions, observer look-fors, and differentiation in a table layout,
    plus anticipated responses and teacher content knowledge sections.
    """
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    from clawed.sanitize import sanitize_text

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    theme = get_color_theme(persona.subject_area or "")
    primary_hex = theme["primary"]
    primary_rgb = RGBColor(int(primary_hex[:2], 16), int(primary_hex[2:4], 16), int(primary_hex[4:6], 16))

    def _shade(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(
            qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color},
        )
        tcPr.append(shading)

    # ── I. Header ─────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_para.add_run("LESSON PLAN")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = primary_rgb

    # Overview table
    teacher_name = getattr(admin_plan, "teacher_name", "") or persona.name or "Teacher"
    overview_data = [
        ("Teacher", sanitize_text(teacher_name)),
        ("Course", sanitize_text(getattr(admin_plan, "course", persona.subject_area or ""))),
        ("Date", sanitize_text(getattr(admin_plan, "date", ""))),
        ("Topic", sanitize_text(getattr(admin_plan, "topic", lesson.title))),
        ("Grade Level", sanitize_text(getattr(admin_plan, "grade_level", ""))),
        ("Duration", f"{getattr(admin_plan, 'duration_minutes', 40)} Minutes"),
        ("Aim", sanitize_text(getattr(admin_plan, "aim", lesson.objective))),
    ]
    standards = getattr(admin_plan, "standards", []) or lesson.standards
    if standards:
        overview_data.append(("Standards", sanitize_text(", ".join(standards[:3]))))
    materials = getattr(admin_plan, "materials", []) or lesson.materials_needed
    if materials:
        overview_data.append(("Materials", sanitize_text(", ".join(materials[:6]))))

    overview_table = doc.add_table(rows=len(overview_data), cols=2)
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(overview_data):
        label_cell = overview_table.rows[i].cells[0]
        value_cell = overview_table.rows[i].cells[1]
        label_cell.text = ""
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        _shade(label_cell, theme.get("bg_light", "F5F5F5").lstrip("#"))
        value_cell.text = value
        for p in value_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
    _set_table_borders(overview_table)

    # ── II. Section-by-Section Breakdown ──────────────────────────────
    sections = getattr(admin_plan, "sections", [])
    if sections:
        doc.add_paragraph("")
        h = doc.add_heading("Lesson Plan & Pacing", level=2)
        for run in h.runs:
            run.font.color.rgb = primary_rgb

        # 5-column table: Section | Teacher Actions | Student Actions | Look-Fors | Differentiation
        col_headers = [
            "Section & Timing", "Teacher Actions", "Student Actions",
            "Observer Look-Fors", "Differentiation",
        ]
        section_table = doc.add_table(rows=1 + len(sections), cols=5)
        section_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for ci, header in enumerate(col_headers):
            cell = section_table.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(header)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(cell, primary_hex)

        # Data rows
        for ri, sec in enumerate(sections):
            row = section_table.rows[ri + 1]
            if hasattr(sec, "section_name"):
                fields = [
                    f"{sec.section_name}\n({sec.timing_minutes} min)",
                    sanitize_text(sec.teacher_actions),
                    sanitize_text(sec.student_actions),
                    sanitize_text(sec.observer_look_fors),
                    sanitize_text(sec.differentiation),
                ]
            else:
                fields = [str(sec)] + [""] * 4
            for ci, text in enumerate(fields):
                cell = row.cells[ci]
                cell.text = text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
            # Alternate row shading
            if ri % 2 == 0:
                for cell in row.cells:
                    _shade(cell, "F9F9F9")

        _set_table_borders(section_table)

    # ── III. Anticipated Responses & Misconceptions ───────────────────
    responses = getattr(admin_plan, "anticipated_responses", [])
    if responses:
        doc.add_paragraph("")
        h = doc.add_heading("Anticipated Student Responses & Misconceptions", level=2)
        for run in h.runs:
            run.font.color.rgb = primary_rgb

        for resp in responses:
            is_mis = getattr(resp, "is_misconception", False) if hasattr(resp, "is_misconception") else False
            text = sanitize_text(getattr(resp, "response_or_misconception", str(resp)))
            correction = sanitize_text(getattr(resp, "teacher_correction", ""))

            p = doc.add_paragraph()
            prefix = "MISCONCEPTION: " if is_mis else "EXPECTED: "
            pr = p.add_run(prefix)
            pr.bold = True
            pr.font.size = Pt(10)
            pr.font.color.rgb = RGBColor(0xCC, 0x33, 0x33) if is_mis else RGBColor(0x33, 0x99, 0x33)
            tr = p.add_run(f'"{text}"')
            tr.italic = True
            tr.font.size = Pt(10)

            if correction:
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.3)
                cr = cp.add_run(f"Redirect: {correction}")
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── IV. Teacher Content Knowledge ─────────────────────────────────
    tck = sanitize_text(getattr(admin_plan, "teacher_content_knowledge", ""))
    if tck:
        doc.add_paragraph("")
        h = doc.add_heading("Teacher Content Knowledge", level=2)
        for run in h.runs:
            run.font.color.rgb = primary_rgb
        p = doc.add_paragraph(tck)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph("")
    gen_footer = doc.add_paragraph()
    gen_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = gen_footer.add_run(f"Generated by {agent_name}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out = _resolve_output(output_dir, lesson, ".docx")
    doc.save(str(out))
    return out
