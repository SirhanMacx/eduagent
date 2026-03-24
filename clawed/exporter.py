"""Export lesson plans and materials to Markdown, PDF, and DOCX formats."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from clawed.models import (
    DailyLesson,
    LessonMaterials,
    PacingGuide,
    UnitPlan,
    YearMap,
)

# ── Markdown export ──────────────────────────────────────────────────────


def unit_to_markdown(unit: UnitPlan) -> str:
    """Render a UnitPlan as formatted Markdown."""
    lines: list[str] = [
        f"# {unit.title}",
        "",
        f"**Subject:** {unit.subject}  ",
        f"**Grade Level:** {unit.grade_level}  ",
        f"**Topic:** {unit.topic}  ",
        f"**Duration:** {unit.duration_weeks} weeks  ",
        "",
        "## Overview",
        "",
        unit.overview,
        "",
        "## Essential Questions",
        "",
    ]
    for q in unit.essential_questions:
        lines.append(f"1. {q}")
    lines.append("")

    lines.append("## Enduring Understandings")
    lines.append("")
    for u in unit.enduring_understandings:
        lines.append(f"- {u}")
    lines.append("")

    if unit.standards:
        lines.append("## Standards")
        lines.append("")
        for s in unit.standards:
            lines.append(f"- {s}")
        lines.append("")

    lines.append("## Daily Lesson Sequence")
    lines.append("")
    lines.append("| # | Topic | Type | Description |")
    lines.append("|---|-------|------|-------------|")
    for lesson in unit.daily_lessons:
        lines.append(
            f"| {lesson.lesson_number} | {lesson.topic} | {lesson.lesson_type} | {lesson.description} |"
        )
    lines.append("")

    lines.append("## Assessment Plan")
    lines.append("")
    lines.append("### Formative")
    for item in unit.assessment_plan.formative:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("### Summative")
    for item in unit.assessment_plan.summative:
        lines.append(f"- {item}")
    lines.append("")

    if unit.required_materials:
        lines.append("## Required Materials")
        lines.append("")
        for m in unit.required_materials:
            lines.append(f"- {m}")
        lines.append("")

    return "\n".join(lines)


def lesson_to_markdown(lesson: DailyLesson) -> str:
    """Render a DailyLesson as formatted Markdown."""
    lines: list[str] = [
        f"# Lesson {lesson.lesson_number}: {lesson.title}",
        "",
        f"**Objective:** {lesson.objective}  ",
    ]

    if lesson.standards:
        lines.append(f"**Standards:** {', '.join(lesson.standards)}  ")
    lines.append("")

    # Time bar
    total = sum(lesson.time_estimates.values())
    lines.append(f"**Total Time:** {total} minutes")
    lines.append("")

    lines.extend([
        "## Do-Now / Warm-Up (5 min)",
        "",
        lesson.do_now,
        "",
        "## Direct Instruction (15-20 min)",
        "",
        lesson.direct_instruction,
        "",
        "## Guided Practice (15-20 min)",
        "",
        lesson.guided_practice,
        "",
        "## Independent Work (10 min)",
        "",
        lesson.independent_work,
        "",
        "## Exit Ticket",
        "",
    ])
    for i, et in enumerate(lesson.exit_ticket, 1):
        lines.append(f"{i}. {et.question}")
        if et.expected_response:
            lines.append(f"   *Expected: {et.expected_response}*")
    lines.append("")

    if lesson.homework:
        lines.extend([
            "## Homework",
            "",
            lesson.homework,
            "",
        ])

    lines.extend([
        "## Differentiation",
        "",
        "### Struggling Learners",
    ])
    for item in lesson.differentiation.struggling:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("### Advanced Learners")
    for item in lesson.differentiation.advanced:
        lines.append(f"- {item}")
    lines.append("")

    if lesson.differentiation.ell:
        lines.append("### ELL Students")
        for item in lesson.differentiation.ell:
            lines.append(f"- {item}")
        lines.append("")

    if lesson.materials_needed:
        lines.append("## Materials Needed")
        lines.append("")
        for m in lesson.materials_needed:
            lines.append(f"- {m}")
        lines.append("")

    return "\n".join(lines)


def materials_to_markdown(materials: LessonMaterials) -> str:
    """Render LessonMaterials as formatted Markdown."""
    lines: list[str] = [
        f"# Materials: {materials.lesson_title}",
        "",
    ]

    # Worksheet
    if materials.worksheet_items:
        lines.extend(["## Student Worksheet", "", "**Name:** ________________________  **Date:** ____________", ""])
        total_pts = sum(item.point_value for item in materials.worksheet_items)
        lines.append(f"**Total Points:** {total_pts}")
        lines.append("")
        for item in materials.worksheet_items:
            pts = "pts" if item.point_value != 1 else "pt"
            lines.append(f"**{item.item_number}.** ({item.point_value} {pts}) {item.prompt}")
            lines.append("")
        lines.append("---")
        lines.append("")

    # Assessment
    if materials.assessment_questions:
        lines.extend(["## Assessment / Quiz", ""])
        total_pts = sum(q.point_value for q in materials.assessment_questions)
        lines.append(f"**Total Points:** {total_pts}")
        lines.append("")
        for q in materials.assessment_questions:
            lines.append(f"**{q.question_number}.** ({q.point_value} pts) {q.question}")
            if q.choices:
                for choice in q.choices:
                    lines.append(f"   {choice}")
            lines.append("")
        lines.append("---")
        lines.append("")

    # Answer Key
    if materials.assessment_questions:
        lines.extend(["## Answer Key", ""])
        for q in materials.assessment_questions:
            lines.append(f"**{q.question_number}.** {q.correct_answer}")
        lines.append("")
        lines.append("---")
        lines.append("")

    # Rubric
    if materials.rubric:
        lines.extend(["## Rubric", ""])
        lines.append("| Criterion | Excellent | Proficient | Developing | Beginning |")
        lines.append("|-----------|-----------|------------|------------|-----------|")
        for r in materials.rubric:
            lines.append(f"| {r.criterion} | {r.excellent} | {r.proficient} | {r.developing} | {r.beginning} |")
        lines.append("")
        lines.append("---")
        lines.append("")

    # Slide Outline
    if materials.slide_outline:
        lines.extend(["## Slide Deck Outline", ""])
        for slide in materials.slide_outline:
            lines.append(f"### Slide {slide.slide_number}: {slide.title}")
            for bullet in slide.content_bullets:
                lines.append(f"- {bullet}")
            if slide.speaker_notes:
                lines.append(f"\n*Speaker Notes: {slide.speaker_notes}*")
            lines.append("")

    # IEP Notes
    if materials.iep_notes:
        lines.extend(["## IEP Accommodation Notes", ""])
        for note in materials.iep_notes:
            lines.append(f"- {note}")
        lines.append("")

    return "\n".join(lines)


# ── Production PDF export (weasyprint) ───────────────────────────────────

_PDF_CSS = """\
@page {
    size: letter;
    margin: 0.75in 1in;
    @top-center { content: "EDUagent Lesson Plan"; font-size: 9pt; color: #888; }
    @bottom-center { content: "Page " counter(page) " of " counter(pages); font-size: 9pt; color: #888; }
}
body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    font-size: 11pt; line-height: 1.6; color: #1a1a2e;
}
h1 { font-size: 20pt; color: #0f3460; margin-bottom: 4pt; page-break-after: avoid; }
h2 {
  font-size: 14pt; color: #16213e; border-bottom: 2px solid #0f3460;
  padding-bottom: 3pt; margin-top: 18pt; page-break-after: avoid;
}
h3 { font-size: 12pt; color: #16213e; margin-top: 12pt; page-break-after: avoid; }
p { margin-bottom: 6pt; }
ul, ol { padding-left: 1.2em; margin-bottom: 8pt; }
li { margin-bottom: 3pt; }
table { width: 100%; border-collapse: collapse; margin-bottom: 12pt; }
th, td { text-align: left; padding: 6pt 8pt; border: 1px solid #ddd; font-size: 10pt; }
th { background: #0f3460; color: #fff; font-weight: 600; }
tr:nth-child(even) { background: #f8f9fa; }
.header-block {
  background: #f0f4ff; border: 1px solid #c7d7f5;
  border-radius: 6px; padding: 12pt 16pt; margin-bottom: 16pt;
}
.header-block p { margin: 2pt 0; font-size: 10pt; }
.section-divider { border: none; border-top: 1px solid #e5e7eb; margin: 12pt 0; }
.worksheet-page { page-break-before: always; }
.score-badge {
  display: inline-block; padding: 2pt 8pt; border-radius: 10pt;
  font-weight: 600; font-size: 10pt; color: #fff;
}
.score-green { background: #059669; }
.score-yellow { background: #d97706; }
.score-red { background: #dc2626; }
"""


def _lesson_to_html_for_pdf(
    lesson: DailyLesson,
    materials: "LessonMaterials | None" = None,
    teacher_name: str = "",
    date_str: str = "",
) -> str:
    """Convert lesson + materials to print-ready HTML for weasyprint."""
    from html import escape as esc

    sections = []

    # Header block
    header_parts = [f"<strong>Objective:</strong> {esc(lesson.objective)}"]
    if lesson.standards:
        header_parts.append(f"<strong>Standards:</strong> {esc(', '.join(lesson.standards))}")
    if teacher_name:
        header_parts.append(f"<strong>Teacher:</strong> {esc(teacher_name)}")
    if date_str:
        header_parts.append(f"<strong>Date:</strong> {esc(date_str)}")
    total_time = sum(lesson.time_estimates.values())
    header_parts.append(f"<strong>Total Time:</strong> {total_time} minutes")
    header_html = "".join(f"<p>{p}</p>" for p in header_parts)

    sections.append(f"<h1>Lesson {lesson.lesson_number}: {esc(lesson.title)}</h1>")
    sections.append(f'<div class="header-block">{header_html}</div>')

    # Lesson sections
    section_map = [
        ("Do-Now / Warm-Up", lesson.do_now),
        ("Direct Instruction", lesson.direct_instruction),
        ("Guided Practice", lesson.guided_practice),
        ("Independent Work", lesson.independent_work),
    ]
    for heading, content in section_map:
        if content:
            sections.append(f"<h2>{heading}</h2><p>{esc(content)}</p>")

    # Exit ticket
    if lesson.exit_ticket:
        sections.append("<h2>Exit Ticket</h2><ol>")
        for et in lesson.exit_ticket:
            sections.append(f"<li>{esc(et.question)}</li>")
        sections.append("</ol>")

    # Homework
    if lesson.homework:
        sections.append(f"<h2>Homework</h2><p>{esc(lesson.homework)}</p>")

    # Differentiation
    diff = lesson.differentiation
    if diff.struggling or diff.advanced or diff.ell:
        sections.append("<h2>Differentiation</h2>")
        if diff.struggling:
            sections.append("<h3>Struggling Learners</h3><ul>")
            for item in diff.struggling:
                sections.append(f"<li>{esc(item)}</li>")
            sections.append("</ul>")
        if diff.advanced:
            sections.append("<h3>Advanced Learners</h3><ul>")
            for item in diff.advanced:
                sections.append(f"<li>{esc(item)}</li>")
            sections.append("</ul>")
        if diff.ell:
            sections.append("<h3>ELL Students</h3><ul>")
            for item in diff.ell:
                sections.append(f"<li>{esc(item)}</li>")
            sections.append("</ul>")

    # Materials on separate page
    if materials and materials.worksheet_items:
        sections.append('<div class="worksheet-page">')
        sections.append(f"<h1>Student Worksheet: {esc(materials.lesson_title)}</h1>")
        sections.append("<p><strong>Name:</strong> ________________________ <strong>Date:</strong> ____________</p>")
        total_pts = sum(item.point_value for item in materials.worksheet_items)
        sections.append(f"<p><strong>Total Points:</strong> {total_pts}</p>")
        for item in materials.worksheet_items:
            pts = "pts" if item.point_value != 1 else "pt"
            prompt_html = esc(item.prompt)
            sections.append(
                f"<p><strong>{item.item_number}.</strong>"
                f" ({item.point_value} {pts}) {prompt_html}</p>"
            )
        sections.append("</div>")

    body = "\n".join(sections)
    return f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{body}</body></html>"


async def export_lesson_pdf(
    lesson: DailyLesson,
    materials: "LessonMaterials | None" = None,
    output_path: Path | None = None,
    teacher_name: str = "",
    date_str: str = "",
) -> Path:
    """Export using weasyprint for professional print-quality PDF."""
    try:
        import weasyprint
    except ImportError:
        # Fall back to reportlab if weasyprint not installed
        md_text = lesson_to_markdown(lesson)
        if output_path is None:
            output_path = Path(f"lesson_{lesson.lesson_number:02d}.pdf")
        return _markdown_to_pdf(md_text, output_path)

    html_content = _lesson_to_html_for_pdf(lesson, materials, teacher_name, date_str)
    if output_path is None:
        output_path = Path(f"lesson_{lesson.lesson_number:02d}.pdf")

    doc = weasyprint.HTML(string=html_content)
    doc.write_pdf(str(output_path), stylesheets=[_pdf_css_stylesheet()])
    return output_path


def _pdf_css_stylesheet():
    """Create a weasyprint CSS stylesheet from our PDF styles."""
    try:
        import weasyprint
        return weasyprint.CSS(string=_PDF_CSS)
    except ImportError:
        return None


# ── PDF export (reportlab fallback) ──────────────────────────────────────


def _markdown_to_pdf(markdown_text: str, output_path: Path) -> Path:
    """Convert Markdown text to a PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    heading2 = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=8,
    )
    body = styles["BodyText"]

    story: list[Any] = []
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], heading1))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], heading2))
        elif stripped.startswith("### "):
            story.append(Paragraph(f"<b>{stripped[4:]}</b>", body))
        elif stripped.startswith("- "):
            story.append(Paragraph(f"&bull; {stripped[2:]}", body))
        elif stripped.startswith("|"):
            # Simplified table rendering — just output the text
            cells = [c.strip() for c in stripped.split("|") if c.strip() and c.strip() != "---"]
            if cells:
                story.append(Paragraph(" | ".join(cells), body))
        elif stripped == "---":
            story.append(Spacer(1, 12))
        else:
            # Handle bold markdown
            text = stripped.replace("**", "<b>", 1).replace("**", "</b>", 1)
            story.append(Paragraph(text, body))

    doc.build(story)
    return output_path


# ── DOCX export ──────────────────────────────────────────────────────────


def _markdown_to_docx(markdown_text: str, output_path: Path) -> Path:
    """Convert Markdown text to a DOCX using python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|") and "---" not in stripped:
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if cells:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(" | ".join(cells))
                run.font.size = Pt(10)
        elif stripped == "---":
            doc.add_paragraph()  # Spacer
        else:
            # Strip markdown bold for plain docx
            clean = stripped.replace("**", "")
            doc.add_paragraph(clean)

    doc.save(str(output_path))
    return output_path


# ── Public API ───────────────────────────────────────────────────────────


def export_unit(unit: UnitPlan, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a unit plan to the specified format."""
    from clawed.io import safe_filename as _safe_fn
    from clawed.io import write_text as _wt

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_fn(unit.title)
    md_text = unit_to_markdown(unit)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        # Graceful fallback — return markdown instead of crashing
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    return path


def export_lesson(lesson: DailyLesson, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a lesson plan to the specified format.

    For pptx/docx/pdf, callers should use doc_export.py instead.
    Unknown formats gracefully fall back to markdown.
    """
    from clawed.io import write_text as _wt

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"lesson_{lesson.lesson_number:02d}"
    md_text = lesson_to_markdown(lesson)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        # Graceful fallback — return markdown instead of crashing
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    return path


def export_materials(materials: LessonMaterials, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export lesson materials to the specified format."""
    from clawed.io import safe_filename as _safe_fn
    from clawed.io import write_text as _wt

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _safe_fn(materials.lesson_title)
    md_text = materials_to_markdown(materials)

    if fmt == "markdown":
        path = output_dir / f"materials_{safe_name}.md"
        _wt(path, md_text)
    elif fmt == "pdf":
        path = output_dir / f"materials_{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"materials_{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        # Graceful fallback — return markdown instead of crashing
        path = output_dir / f"materials_{safe_name}.md"
        _wt(path, md_text)
    return path


# ── Year map / pacing exports ────────────────────────────────────────────


def year_map_to_markdown(year_map: YearMap) -> str:
    """Render a YearMap as formatted Markdown."""
    lines: list[str] = [
        f"# Full-Year Curriculum Map — {year_map.subject}, Grade {year_map.grade_level}",
        "",
        f"**School Year:** {year_map.school_year or 'N/A'}  ",
        f"**Total Instructional Weeks:** {year_map.total_weeks}",
        "",
    ]

    # Units table
    lines.append("## Unit Sequence")
    lines.append("")
    lines.append("| # | Unit | Weeks | Standards |")
    lines.append("|---|------|-------|-----------|")
    for u in year_map.units:
        stds = ", ".join(u.standards[:3])
        if len(u.standards) > 3:
            stds += f" (+{len(u.standards) - 3} more)"
        lines.append(f"| {u.unit_number} | {u.title} | {u.duration_weeks} | {stds} |")
    lines.append("")

    # Unit details
    for u in year_map.units:
        lines.append(f"### Unit {u.unit_number}: {u.title}")
        lines.append(f"**Duration:** {u.duration_weeks} weeks  ")
        if u.description:
            lines.append(f"\n{u.description}")
        if u.essential_questions:
            lines.append("\n**Essential Questions:**")
            for eq in u.essential_questions:
                lines.append(f"- {eq}")
        if u.standards:
            lines.append(f"\n**Standards:** {', '.join(u.standards)}")
        lines.append("")

    # Big ideas
    if year_map.big_ideas:
        lines.append("## Big Ideas (Cross-Unit Connections)")
        lines.append("")
        for bi in year_map.big_ideas:
            unit_refs = ", ".join(str(n) for n in bi.connected_units)
            lines.append(f"- **{bi.idea}** (Units {unit_refs})")
        lines.append("")

    # Assessment calendar
    if year_map.assessment_calendar:
        lines.append("## Assessment Calendar")
        lines.append("")
        lines.append("| Week | Type | Assessment | Unit |")
        lines.append("|------|------|------------|------|")
        for a in sorted(year_map.assessment_calendar, key=lambda x: x.week):
            unit_label = f"Unit {a.unit_number}" if a.unit_number > 0 else "—"
            lines.append(f"| {a.week} | {a.assessment_type.title()} | {a.title} | {unit_label} |")
        lines.append("")

    return "\n".join(lines)


def pacing_guide_to_markdown(guide: PacingGuide) -> str:
    """Render a PacingGuide as formatted Markdown."""
    lines: list[str] = [
        f"# Pacing Guide — {guide.subject}, Grade {guide.grade_level}",
        "",
        f"**School Year:** {guide.school_year or 'N/A'}  ",
        f"**Start Date:** {guide.start_date}",
        "",
        "| Week | Dates | Unit | Topics | Notes |",
        "|------|-------|------|--------|-------|",
    ]

    for w in guide.weeks:
        topics = "; ".join(w.topics) if w.topics else "—"
        notes = w.notes or ""
        lines.append(
            f"| {w.week_number} | {w.start_date} – {w.end_date} | "
            f"U{w.unit_number}: {w.unit_title} | {topics} | {notes} |"
        )

    lines.append("")
    return "\n".join(lines)


def export_year_map(year_map: YearMap, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a year map to the specified format."""
    from clawed.io import safe_filename as _safe_fn
    from clawed.io import write_text as _wt

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"year_map_{_safe_fn(year_map.subject)}_{year_map.grade_level}"
    md_text = year_map_to_markdown(year_map)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    return path


def export_pacing_guide(guide: PacingGuide, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a pacing guide to the specified format."""
    from clawed.io import safe_filename as _safe_fn
    from clawed.io import write_text as _wt

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"pacing_{_safe_fn(guide.subject)}_{guide.grade_level}"
    md_text = pacing_guide_to_markdown(guide)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        path = output_dir / f"{safe_name}.md"
        _wt(path, md_text)
    return path
