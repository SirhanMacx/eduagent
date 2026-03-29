"""Tests for mechanical compilation of teacher/student/slide views from MasterContent."""

from __future__ import annotations

import asyncio
from pathlib import Path

import pytest

from clawed.master_content import (
    DoNow,
    GuidedNote,
    IndependentWork,
    InstructionSection,
    MasterContent,
    PrimarySource,
    StationDocument,
    StimulusQuestion,
    VocabularyEntry,
)
from clawed.models import DifferentiationNotes


# ── Shared fixture ──────────────────────────────────────────────────────────


def _make_master() -> MasterContent:
    """Return a realistic MasterContent for compilation tests."""
    return MasterContent(
        title="The Industrial Revolution",
        subject="Social Studies",
        grade_level="8",
        topic="Industrial Revolution",
        standards=["SS.8.1", "SS.8.2"],
        objective="Students will analyze causes of the Industrial Revolution.",
        duration_minutes=45,
        vocabulary=[
            VocabularyEntry(
                term="industrialization",
                definition="Process of developing industry on a large scale.",
                context_sentence="Industrialization transformed daily life in 19th-century Britain.",
            ),
            VocabularyEntry(
                term="urbanization",
                definition="Rapid growth of towns and cities.",
                context_sentence="Urbanization followed the shift from farm to factory work.",
            ),
        ],
        primary_sources=[
            PrimarySource(
                id="ps1",
                title="Factory Conditions Report, 1833",
                source_type="text_excerpt",
                content_text="Children as young as six worked 14-hour days in dangerous conditions.",
                attribution="Parliamentary Commission, 1833",
                scaffolding_questions=[
                    "What working conditions are described?",
                    "Who is the intended audience for this report?",
                ],
            )
        ],
        do_now=DoNow(
            stimulus="Look at the image of a 19th-century factory.",
            stimulus_type="image",
            questions=["What do you notice?", "What do you wonder?"],
            answers=["Students observe machinery and workers.", "Responses vary."],
        ),
        direct_instruction=[
            InstructionSection(
                heading="Background: Pre-Industrial Britain",
                content="Before 1750, most goods were made by hand in rural homes.",
                teacher_script="Say: 'Imagine a world without machines…'",
                key_points=["Cottage industry", "Agricultural economy"],
            ),
            InstructionSection(
                heading="Rise of the Factory System",
                content="Steam-powered machinery moved production into large factories.",
                teacher_script="Ask: 'How would your daily life change?'",
                key_points=["Steam engine", "Mass production", "Child labour"],
            ),
        ],
        guided_notes=[
            GuidedNote(
                prompt="The Industrial Revolution began in ______.",
                answer="ANSWER_BRITAIN_UNIQUE",
                section_ref="Background: Pre-Industrial Britain",
            ),
            GuidedNote(
                prompt="The steam engine was invented by ______.",
                answer="ANSWER_WATT_UNIQUE",
                section_ref="Rise of the Factory System",
            ),
        ],
        stations=[
            StationDocument(
                title="Station A: Factory Life",
                source_ref="ps1",
                task="Read the excerpt and answer the questions.",
                student_directions="Read carefully, then respond in complete sentences.",
                teacher_answer_key="Children worked long hours in dangerous conditions.",
            )
        ],
        independent_work=IndependentWork(
            task="Write a paragraph comparing factory and farm work.",
            rubric_snippet="4 = strong evidence; 1 = no evidence",
        ),
        exit_ticket=[
            StimulusQuestion(
                stimulus="Textile workers earned less than 1 shilling per day.",
                stimulus_type="text_excerpt",
                question="How did factory wages affect working-class families?",
                answer="Low wages forced multiple family members to work.",
                cognitive_level="analysis",
            )
        ],
        differentiation=DifferentiationNotes(
            struggling=["Provide sentence starters"],
            advanced=["Compare to today's globalization"],
            ell=["Provide bilingual vocabulary sheet"],
        ),
        homework="Read pages 45–52 and answer questions 1–3.",
        materials_needed=["Chromebook", "Student packet"],
    )


# ── Helper ──────────────────────────────────────────────────────────────────


def _run(coro):
    """Run a coroutine synchronously (tests are not async)."""
    return asyncio.run(coro)


# ── test 1: teacher view creates a non-empty DOCX ──────────────────────────


def test_teacher_view_creates_docx(tmp_path):
    from clawed.compile_teacher import compile_teacher_view

    mc = _make_master()
    out = _run(compile_teacher_view(mc, images={}, output_dir=tmp_path))

    assert out.exists(), f"Expected file at {out}"
    assert out.suffix == ".docx"
    assert out.stat().st_size > 0, "DOCX must be non-empty"


# ── test 2: student view creates a non-empty DOCX ─────────────────────────


def test_student_view_creates_docx(tmp_path):
    from clawed.compile_student import compile_student_view

    mc = _make_master()
    out = _run(compile_student_view(mc, images={}, output_dir=tmp_path))

    assert out.exists(), f"Expected file at {out}"
    assert out.suffix == ".docx"
    assert out.stat().st_size > 0, "DOCX must be non-empty"


# ── test 3: slides creates a non-empty PPTX ───────────────────────────────


def test_slides_creates_pptx(tmp_path):
    from clawed.compile_slides import compile_slides

    mc = _make_master()
    out = _run(compile_slides(mc, images={}, output_dir=tmp_path))

    assert out.exists(), f"Expected file at {out}"
    assert out.suffix == ".pptx"
    assert out.stat().st_size > 0, "PPTX must be non-empty"


# ── test 4: student view has blanks, not answers ──────────────────────────


def test_student_view_has_blanks(tmp_path):
    from docx import Document

    from clawed.compile_student import compile_student_view

    mc = _make_master()
    out = _run(compile_student_view(mc, images={}, output_dir=tmp_path))

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Blank lines should be present in guided notes section
    assert "_____________" in full_text, "Student view must contain blank lines for guided notes"

    # The actual guided note answers must NOT appear
    assert "ANSWER_BRITAIN_UNIQUE" not in full_text, (
        "Guided note answer must not appear in student view"
    )
    assert "ANSWER_WATT_UNIQUE" not in full_text, (
        "Guided note answer must not appear in student view"
    )


# ── test 5: teacher view has answers filled in ────────────────────────────


def test_teacher_view_has_answers(tmp_path):
    from docx import Document

    from clawed.compile_teacher import compile_teacher_view

    mc = _make_master()
    out = _run(compile_teacher_view(mc, images={}, output_dir=tmp_path))

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Both guided note answers must appear
    assert "ANSWER_BRITAIN_UNIQUE" in full_text, (
        "Teacher view must contain guided note answer"
    )
    assert "ANSWER_WATT_UNIQUE" in full_text, (
        "Teacher view must contain guided note answer"
    )

    # Teacher script must appear
    assert "Imagine a world without machines" in full_text, (
        "Teacher view must include teacher scripts"
    )

    # Station answer key must appear
    assert "Children worked long hours" in full_text, (
        "Teacher view must include station answer keys"
    )


# ── test 6: slide count ───────────────────────────────────────────────────


def test_slide_count(tmp_path):
    from pptx import Presentation

    from clawed.compile_slides import compile_slides

    mc = _make_master()
    out = _run(compile_slides(mc, images={}, output_dir=tmp_path))

    prs = Presentation(str(out))
    slide_count = len(prs.slides)

    # Expected:
    #   1 title slide
    # + 1 vocabulary slide (2 terms, fits on one)
    # + 2 instruction section slides
    # + 1 source analysis slide (1 primary source)
    # + 1 station overview slide
    # + 1 exit ticket slide
    # = 7 total
    expected = (
        1  # title
        + 1  # vocabulary (2 terms on one slide, threshold is 5)
        + len(mc.direct_instruction)       # 2
        + len(mc.primary_sources)          # 1
        + (1 if mc.stations else 0)        # 1
        + (1 if mc.exit_ticket else 0)     # 1
    )
    assert slide_count == expected, (
        f"Expected {expected} slides, got {slide_count}"
    )
