"""Export lesson plans and materials to Markdown, PDF, and DOCX formats."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from eduagent.models import (
    DailyLesson,
    LessonMaterials,
    UnitPlan,
)


# ── Markdown export ──────────────────────────────────────────────────────


def unit_to_markdown(unit: UnitPlan) -> str:
    """Render a UnitPlan as formatted Markdown."""
    lines: list[str] = [
        f"# {unit.title}",
        "",
        f"**Subject:** {unit.subject}  ",
        f"**Grade Level:** {unit.grade_level}  ",
        f"**Topic:** {unit.topic}  ",
        f"**Duration:** {unit.duration_weeks} weeks  ",
        "",
        "## Overview",
        "",
        unit.overview,
        "",
        "## Essential Questions",
        "",
    ]
    for q in unit.essential_questions:
        lines.append(f"1. {q}")
    lines.append("")

    lines.append("## Enduring Understandings")
    lines.append("")
    for u in unit.enduring_understandings:
        lines.append(f"- {u}")
    lines.append("")

    if unit.standards:
        lines.append("## Standards")
        lines.append("")
        for s in unit.standards:
            lines.append(f"- {s}")
        lines.append("")

    lines.append("## Daily Lesson Sequence")
    lines.append("")
    lines.append("| # | Topic | Type | Description |")
    lines.append("|---|-------|------|-------------|")
    for lesson in unit.daily_lessons:
        lines.append(
            f"| {lesson.lesson_number} | {lesson.topic} | {lesson.lesson_type} | {lesson.description} |"
        )
    lines.append("")

    lines.append("## Assessment Plan")
    lines.append("")
    lines.append("### Formative")
    for item in unit.assessment_plan.formative:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("### Summative")
    for item in unit.assessment_plan.summative:
        lines.append(f"- {item}")
    lines.append("")

    if unit.required_materials:
        lines.append("## Required Materials")
        lines.append("")
        for m in unit.required_materials:
            lines.append(f"- {m}")
        lines.append("")

    return "\n".join(lines)


def lesson_to_markdown(lesson: DailyLesson) -> str:
    """Render a DailyLesson as formatted Markdown."""
    lines: list[str] = [
        f"# Lesson {lesson.lesson_number}: {lesson.title}",
        "",
        f"**Objective:** {lesson.objective}  ",
    ]

    if lesson.standards:
        lines.append(f"**Standards:** {', '.join(lesson.standards)}  ")
    lines.append("")

    # Time bar
    total = sum(lesson.time_estimates.values())
    lines.append(f"**Total Time:** {total} minutes")
    lines.append("")

    lines.extend([
        "## Do-Now / Warm-Up (5 min)",
        "",
        lesson.do_now,
        "",
        "## Direct Instruction (15-20 min)",
        "",
        lesson.direct_instruction,
        "",
        "## Guided Practice (15-20 min)",
        "",
        lesson.guided_practice,
        "",
        "## Independent Work (10 min)",
        "",
        lesson.independent_work,
        "",
        "## Exit Ticket",
        "",
    ])
    for i, et in enumerate(lesson.exit_ticket, 1):
        lines.append(f"{i}. {et.question}")
        if et.expected_response:
            lines.append(f"   *Expected: {et.expected_response}*")
    lines.append("")

    if lesson.homework:
        lines.extend([
            "## Homework",
            "",
            lesson.homework,
            "",
        ])

    lines.extend([
        "## Differentiation",
        "",
        "### Struggling Learners",
    ])
    for item in lesson.differentiation.struggling:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("### Advanced Learners")
    for item in lesson.differentiation.advanced:
        lines.append(f"- {item}")
    lines.append("")

    if lesson.differentiation.ell:
        lines.append("### ELL Students")
        for item in lesson.differentiation.ell:
            lines.append(f"- {item}")
        lines.append("")

    if lesson.materials_needed:
        lines.append("## Materials Needed")
        lines.append("")
        for m in lesson.materials_needed:
            lines.append(f"- {m}")
        lines.append("")

    return "\n".join(lines)


def materials_to_markdown(materials: LessonMaterials) -> str:
    """Render LessonMaterials as formatted Markdown."""
    lines: list[str] = [
        f"# Materials: {materials.lesson_title}",
        "",
    ]

    # Worksheet
    if materials.worksheet_items:
        lines.extend(["## Student Worksheet", "", f"**Name:** ________________________  **Date:** ____________", ""])
        total_pts = sum(item.point_value for item in materials.worksheet_items)
        lines.append(f"**Total Points:** {total_pts}")
        lines.append("")
        for item in materials.worksheet_items:
            lines.append(f"**{item.item_number}.** ({item.point_value} pt{'s' if item.point_value != 1 else ''}) {item.prompt}")
            lines.append("")
        lines.append("---")
        lines.append("")

    # Assessment
    if materials.assessment_questions:
        lines.extend(["## Assessment / Quiz", ""])
        total_pts = sum(q.point_value for q in materials.assessment_questions)
        lines.append(f"**Total Points:** {total_pts}")
        lines.append("")
        for q in materials.assessment_questions:
            lines.append(f"**{q.question_number}.** ({q.point_value} pts) {q.question}")
            if q.choices:
                for choice in q.choices:
                    lines.append(f"   {choice}")
            lines.append("")
        lines.append("---")
        lines.append("")

    # Answer Key
    if materials.assessment_questions:
        lines.extend(["## Answer Key", ""])
        for q in materials.assessment_questions:
            lines.append(f"**{q.question_number}.** {q.correct_answer}")
        lines.append("")
        lines.append("---")
        lines.append("")

    # Rubric
    if materials.rubric:
        lines.extend(["## Rubric", ""])
        lines.append("| Criterion | Excellent | Proficient | Developing | Beginning |")
        lines.append("|-----------|-----------|------------|------------|-----------|")
        for r in materials.rubric:
            lines.append(f"| {r.criterion} | {r.excellent} | {r.proficient} | {r.developing} | {r.beginning} |")
        lines.append("")
        lines.append("---")
        lines.append("")

    # Slide Outline
    if materials.slide_outline:
        lines.extend(["## Slide Deck Outline", ""])
        for slide in materials.slide_outline:
            lines.append(f"### Slide {slide.slide_number}: {slide.title}")
            for bullet in slide.content_bullets:
                lines.append(f"- {bullet}")
            if slide.speaker_notes:
                lines.append(f"\n*Speaker Notes: {slide.speaker_notes}*")
            lines.append("")

    # IEP Notes
    if materials.iep_notes:
        lines.extend(["## IEP Accommodation Notes", ""])
        for note in materials.iep_notes:
            lines.append(f"- {note}")
        lines.append("")

    return "\n".join(lines)


# ── PDF export ───────────────────────────────────────────────────────────


def _markdown_to_pdf(markdown_text: str, output_path: Path) -> Path:
    """Convert Markdown text to a PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    heading2 = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=8,
    )
    body = styles["BodyText"]

    story: list[Any] = []
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], heading1))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], heading2))
        elif stripped.startswith("### "):
            story.append(Paragraph(f"<b>{stripped[4:]}</b>", body))
        elif stripped.startswith("- "):
            story.append(Paragraph(f"&bull; {stripped[2:]}", body))
        elif stripped.startswith("|"):
            # Simplified table rendering — just output the text
            cells = [c.strip() for c in stripped.split("|") if c.strip() and c.strip() != "---"]
            if cells:
                story.append(Paragraph(" | ".join(cells), body))
        elif stripped == "---":
            story.append(Spacer(1, 12))
        else:
            # Handle bold markdown
            text = stripped.replace("**", "<b>", 1).replace("**", "</b>", 1)
            story.append(Paragraph(text, body))

    doc.build(story)
    return output_path


# ── DOCX export ──────────────────────────────────────────────────────────


def _markdown_to_docx(markdown_text: str, output_path: Path) -> Path:
    """Convert Markdown text to a DOCX using python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|") and "---" not in stripped:
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if cells:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(" | ".join(cells))
                run.font.size = Pt(10)
        elif stripped == "---":
            doc.add_paragraph()  # Spacer
        else:
            # Strip markdown bold for plain docx
            clean = stripped.replace("**", "")
            doc.add_paragraph(clean)

    doc.save(str(output_path))
    return output_path


# ── Public API ───────────────────────────────────────────────────────────


def export_unit(unit: UnitPlan, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a unit plan to the specified format."""
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = unit.title.lower().replace(" ", "_")[:50]
    md_text = unit_to_markdown(unit)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        path.write_text(md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return path


def export_lesson(lesson: DailyLesson, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export a lesson plan to the specified format."""
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"lesson_{lesson.lesson_number:02d}"
    md_text = lesson_to_markdown(lesson)

    if fmt == "markdown":
        path = output_dir / f"{safe_name}.md"
        path.write_text(md_text)
    elif fmt == "pdf":
        path = output_dir / f"{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return path


def export_materials(materials: LessonMaterials, output_dir: Path, fmt: str = "markdown") -> Path:
    """Export lesson materials to the specified format."""
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = materials.lesson_title.lower().replace(" ", "_")[:50]
    md_text = materials_to_markdown(materials)

    if fmt == "markdown":
        path = output_dir / f"materials_{safe_name}.md"
        path.write_text(md_text)
    elif fmt == "pdf":
        path = output_dir / f"materials_{safe_name}.pdf"
        _markdown_to_pdf(md_text, path)
    elif fmt == "docx":
        path = output_dir / f"materials_{safe_name}.docx"
        _markdown_to_docx(md_text, path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return path
